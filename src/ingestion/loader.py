from pathlib import Path
from dataclasses import dataclass
from openpyxl import load_workbook
import csv

@dataclass
class Document:
    text: str
    metadata: dict


def load_pdf(path: Path) -> list[Document]:
    import pymupdf

    docs = []
    with pymupdf.open(path) as pdf:
        for page_num, page in enumerate(pdf, start=1):
            text = page.get_text().strip()
            if text:
                docs.append(Document(
                    text=text,
                    metadata={"source": path.name, "page": page_num},
                ))
    return docs


def load_docx(path: Path) -> list[Document]:
    from docx import Document as DocxDocument

    doc = DocxDocument(path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if text:
        return [Document(text=text, metadata={"source": path.name})]
    return []


def load_txt(path: Path) -> list[Document]:
    text = path.read_text(encoding="utf-8").strip()
    if text:
        return [Document(text=text, metadata={"source": path.name})]
    return []

def load_csv(path: Path) -> list[Document]:
    docs = []
    with open(path, newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        rows = list(reader)

    if not rows:
        return []

    # Small files: single document with full table
    if len(rows) <= 50:
        lines = [", ".join(f"{k}: {v}" for k, v in row.items() if v) for row in rows]
        text = "\n".join(lines)
        return [Document(text=text, metadata={"source": path.name})]

    # Large files: group into batches of 20 rows
    for i in range(0, len(rows), 20):
        batch = rows[i:i + 20]
        lines = [", ".join(f"{k}: {v}" for k, v in row.items() if v) for row in batch]
        text = "\n".join(lines)
        docs.append(Document(
            text=text,
            metadata={"source": path.name, "rows": f"{i + 1}-{i + len(batch)}"},
        ))
    return docs

def load_excel(path: Path) -> list[Document]:
    wb = load_workbook(path, read_only=True, data_only=True)
    docs = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))

        if not rows:
            continue

        # First row as headers
        headers = [str(h) if h else f"col_{i}" for i, h in enumerate(rows[0])]
        data_rows = rows[1:]

        if not data_rows:
            continue

        # Same logic: small sheets as one doc, large ones batched
        if len(data_rows) <= 50:
            lines = [
                ", ".join(f"{headers[i]}: {v}" for i, v in enumerate(row) if v is not None)
                for row in data_rows
            ]
            text = "\n".join(lines)
            docs.append(Document(
                text=text,
                metadata={"source": path.name, "sheet": sheet_name},
            ))
        else:
            for batch_start in range(0, len(data_rows), 20):
                batch = data_rows[batch_start:batch_start + 20]
                lines = [
                    ", ".join(f"{headers[i]}: {v}" for i, v in enumerate(row) if v is not None)
                    for row in batch
                ]
                text = "\n".join(lines)
                docs.append(Document(
                    text=text,
                    metadata={
                        "source": path.name,
                        "sheet": sheet_name,
                        "rows": f"{batch_start + 2}-{batch_start + 1 + len(batch)}",
                    },
                ))

    wb.close()
    return docs

LOADERS = {
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".txt": load_txt,
    ".csv": load_csv,
    ".xlsx": load_excel,
    ".xls": load_excel,
}


def load_document(path: Path) -> list[Document]:
    suffix = path.suffix.lower()
    loader = LOADERS.get(suffix)
    if not loader:
        raise ValueError(f"Unsupported file type: {suffix}")
    return loader(path)
